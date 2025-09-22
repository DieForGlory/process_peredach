# /app/cadastre_process/workflows/group_1_workflow.py
import docx
from io import BytesIO

def generate_unilateral_act(deal_data: dict):
    """
    Генерирует Word-документ одностороннего акта.
    """
    # В будущем сюда можно передавать больше данных о сделке
    client_name = deal_data.get('client_name', 'Клиент')
    apartment_id = deal_data.get('property_id', 'N/A')

    doc = docx.Document()
    doc.add_heading('ОДНОСТОРОННИЙ АКТ ПРИЕМА-ПЕРЕДАЧИ', 0)
    p = doc.add_paragraph()
    p.add_run(f'Настоящий акт составлен в связи с неявкой клиента ({client_name}) ').bold = True
    p.add_run('в установленный 30-дневный срок для приёмки квартиры №')
    p.add_run(str(apartment_id)).bold = True
    p.add_run('.')
    # ... здесь можно добавить больше текста и деталей акта ...

    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer