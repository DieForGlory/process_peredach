# /app/cadastre_process/workflows/group_2_workflow.py
import io
import docx
from datetime import datetime, timedelta
from app import db
from ..models import DealStatus
from ..services.data_service import get_single_deal_details_for_workflow, check_debt_status_from_mysql
from flask import current_app


# --- Точка входа для планировщика ---

def check_all_debt_statuses():
    """
    Главная функция, вызываемая планировщиком (н-р, каждые 3 часа).
    Проверяет все сделки в статусах, требующих мониторинга долга.
    """
    with current_app.app_context():
        deals_to_check = DealStatus.query.filter(
            DealStatus.status.in_([
                'pending_debt_payment',
                'penalty_accrual'
            ])
        ).all()

        current_time = datetime.utcnow()

        for deal_status in deals_to_check:
            # 1. Получаем АКТУАЛЬНЫЙ статус долга из MySQL
            fresh_has_debt = check_debt_status_from_mysql(deal_status.deal_id)

            if deal_status.status == 'pending_debt_payment':
                _handle_pending_debt(deal_status, fresh_has_debt, current_time)

            elif deal_status.status == 'penalty_accrual':
                _handle_penalty_accrual(deal_status, fresh_has_debt, current_time)

        try:
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            print(f"Ошибка при пакетном обновлении статусов должников: {e}")


# --- Внутренняя логика ---

def _handle_pending_debt(deal_status, has_debt, current_time):
    """Обрабатывает сделки на 10-дневном таймере."""

    if not has_debt:
        # Долг погашен! Переводим на сценарий Группы 1
        deal_status.status = 'pending_arrival'
        deal_status.documents_delivered_at = current_time  # Запускаем 30-дневный таймер
        deal_status.debt_payment_deadline = None

    elif deal_status.debt_payment_deadline and current_time > deal_status.debt_payment_deadline:
        # 10 дней прошло, долг не погашен. Начинаем расчет пени.
        deal_status.status = 'penalty_accrual'
        deal_status.penalty_check_deadline = current_time + timedelta(days=2)  # Запускаем 2-дневный таймер
        _calculate_and_notify_penalty(deal_status)


def _handle_penalty_accrual(deal_status, has_debt, current_time):
    """Обрабатывает сделки на 2-дневном таймере (начисление пени)."""

    if not has_debt:
        # Долг погашен! Переводим на сценарий Группы 1
        deal_status.status = 'pending_arrival'
        deal_status.documents_delivered_at = current_time  # Запускаем 30-дневный таймер
        deal_status.penalty_check_deadline = None

    elif deal_status.penalty_check_deadline and current_time > deal_status.penalty_check_deadline:
        # 2 дня прошло, долг все еще не погашен.
        deal_status.status = 'termination_pending'
        _send_termination_email(deal_status)


def _calculate_and_notify_penalty(deal_status):
    """
    Расчет пени и запуск уведомлений, если превышен лимит.
    """
    # Получаем фин. данные (сумма, дата) из MySQL
    financials = get_single_deal_details_for_workflow(deal_status.deal_id)

    contract_sum = financials.get('deal_sum', 0)

    # ИСПОЛЬЗУЕМ ОБНОВЛЕННОЕ ПОЛЕ (самая ранняя просрочка)
    due_date = financials.get('first_overdue_payment_date')

    if not due_date:
        # Если даты нет (редкий случай), считаем просрочку от дедлайна (10 дней назад)
        due_date = deal_status.debt_payment_deadline - timedelta(days=10)

    days_overdue = (datetime.utcnow() - due_date).days

    if days_overdue <= 0 or contract_sum == 0:
        # Некорректные данные, пропускаем расчет
        return

    # 0.01% от суммы договора за каждый день
    penalty = days_overdue * (contract_sum * 0.0001)
    deal_status.current_penalty_amount = penalty

    # Проверяем лимит (3% от суммы договора)
    if penalty > (contract_sum * 0.03):
        # Генерируем документ (в этой реализации просто сохраняем сумму)
        deal_status.penalty_notification_generated = True

        # Отправляем email сотруднику
        _send_employee_email(deal_status, penalty, contract_sum)


# --- Генерация документов и E-mail (Заглушки) ---

def get_penalty_notification_doc(deal_id):
    """
    Генерирует (на лету) или берет сохраненный документ уведомления о пене.
    """
    deal_status = DealStatus.query.get(deal_id)
    if not deal_status or not deal_status.current_penalty_amount:
        return None

    financials = get_single_deal_details_for_workflow(deal_id)

    doc = docx.Document()
    doc.add_heading('УВЕДОМЛЕНИЕ О ПЕНЕ', 0)
    p = doc.add_paragraph(f"Уважаемый(ая) {financials.get('client_name', 'Клиент')},")
    p.add_run(f"\n\nПо Вашей квартире №{financials.get('property_id', 'N/A')} имеется просроченная задолженность.")
    p.add_run(f"\nСумма договора: {financials.get('deal_sum', 0):.2f} у.е.")
    p.add_run(f"\nНа текущий момент ({datetime.utcnow().strftime('%Y-%m-%d')}) сумма начисленной пени составляет: ")
    p.add_run(f"{deal_status.current_penalty_amount:.2f} у.е.").bold = True
    p.add_run(f"\nДанная сумма превысила 3% от суммы договора.")
    p.add_run("\n\nПросим Вас срочно погасить задолженность.")

    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer, financials.get('property_id', 'doc')


def _send_employee_email(deal_status, penalty, contract_sum):
    """(Заглушка) Отправляет email ответственному сотруднику."""
    print("--- ОТПРАВКА EMAIL СОТРУДНИКУ ---")
    print(f"ID Сделки: {deal_status.deal_id}")
    print(f"Сумма пени: {penalty:.2f} (превысила 3% от {contract_sum:.2f})")
    # TODO: Добавить реальную логику отправки email
    # email_service.send_mail(
    #     to='responsible_manager@example.com',
    #     subject=f'ПЕНЯ ПРЕВЫСИЛА ЛИМИТ: Сделка {deal_status.deal_id}',
    #     body=f'...'
    # )
    print("-----------------------------------")


def _send_termination_email(deal_status):
    """(Заглушка) Отправляет email о расторжении."""
    print("--- ОТПРАВКА EMAIL О РАСТОРЖЕНИИ ---")
    print(f"ID Сделки: {deal_status.deal_id}")
    print("Клиент не погасил долг в 2-дневный срок после начисления пени.")
    # TODO: Добавить реальную логику отправки email
    # email_service.send_mail(
    #     to='legal_department@example.com, management@example.com',
    #     subject=f'ИНИЦИАЦИЯ РАСТОРЖЕНИЯ: Сделка {deal_status.deal_id}',
    #     body=f'...'
    # )
    print("---------------------------------------")