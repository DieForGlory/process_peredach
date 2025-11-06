# /app/cadastre_process/workflows/group_5_workflow.py
import io
import docx
from datetime import datetime, timedelta
from app import db
from ..models import DealStatus
from ..services.data_service import (
    get_single_deal_details_for_workflow,
    check_increase_payment
)
from flask import current_app


# --- Точка входа для планировщика ---

def check_all_increase_statuses():
    """
    Главная функция, вызываемая планировщиком.
    Проверяет все сделки в статусах, требующих мониторинга (Группы 3 и 5).
    """
    with current_app.app_context():
        deals_to_check = DealStatus.query.filter(
            DealStatus.status.in_([
                'pending_increase_signing',
                'pending_increase_payment',
                'pending_increase_penalty'
            ])
        ).all()

        current_time = datetime.utcnow()

        for deal_status in deals_to_check:
            if deal_status.status == 'pending_increase_signing':
                _handle_pending_signing(deal_status, current_time)

            elif deal_status.status == 'pending_increase_payment':
                _handle_pending_payment(deal_status, current_time)

            elif deal_status.status == 'pending_increase_penalty':
                _handle_pending_penalty(deal_status, current_time)

        try:
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            print(f"Ошибка при пакетном обновлении статусов (Увеличение площади): {e}")


# --- Внутренняя логика ---

def _handle_pending_signing(deal_status, current_time):
    """
    Обрабатывает сделки на 10-дневном таймере (ожидание подписания ДС).
    Если 10 дней прошло -> начисляем штраф (Path B).
    """
    if deal_status.area_increase_agreement_deadline and current_time > deal_status.area_increase_agreement_deadline:
        # 10 дней прошло, ДС не подписано. Начинаем Path B.
        financials = get_single_deal_details_for_workflow(deal_status.deal_id)
        contract_sum = financials.get('deal_sum', 0)
        surcharge_amount = deal_status.area_increase_payment_amount or 0

        # Штраф = 10% от договора + Сумма доплаты
        penalty_amount = (contract_sum * 0.10) + surcharge_amount

        deal_status.status = 'pending_increase_penalty'
        deal_status.area_increase_penalty_amount = penalty_amount
        # Запускаем 15-дневный таймер на оплату штрафа
        deal_status.area_increase_penalty_deadline = current_time + timedelta(days=15)
        deal_status.area_increase_penalty_doc_generated = True  # Флаг для кнопки

        # Сбрасываем 10-дневный таймер
        deal_status.area_increase_agreement_deadline = None


def _handle_pending_payment(deal_status, current_time):
    """
    Обрабатывает сделки на 30-дневном таймере (ожидание оплаты ДС).
    Path A.
    """
    required_amount = deal_status.area_increase_payment_amount or 0
    if required_amount == 0:
        return  # Ошибка, сумма не была рассчитана

    # Проверяем, поступил ли платеж
    payment_found = check_increase_payment(deal_status.deal_id, required_amount)

    if payment_found:
        # Оплата найдена! Переводим на Группу 1 (ожидание визита)
        deal_status.status = 'pending_arrival'
        deal_status.documents_delivered_at = current_time  # Запускаем 30-дневный таймер визита
        deal_status.area_increase_payment_deadline = None

    elif deal_status.area_increase_payment_deadline and current_time > deal_status.area_increase_payment_deadline:
        # 30 дней прошло, оплаты нет.
        deal_status.status = 'termination_pending'
        _send_termination_email(deal_status, "Неуплата доп. соглашения (30 дней)")


def _handle_pending_penalty(deal_status, current_time):
    """
    Обрабатывает сделки на 15-дневном таймере (ожидание оплаты ШТРАФА).
    Path B.
    """
    required_amount = deal_status.area_increase_penalty_amount or 0
    if required_amount == 0:
        return

    # Проверяем, поступил ли платеж
    payment_found = check_increase_payment(deal_status.deal_id, required_amount)

    if payment_found:
        # Штраф оплачен! Переводим на Группу 1 (ожидание визита)
        deal_status.status = 'pending_arrival'
        deal_status.documents_delivered_at = current_time  # Запускаем 30-дневный таймер визита
        deal_status.area_increase_penalty_deadline = None

    elif deal_status.area_increase_penalty_deadline and current_time > deal_status.area_increase_penalty_deadline:
        # 15 дней прошло, оплаты штрафа нет.
        deal_status.status = 'termination_pending'
        _send_termination_email(deal_status, "Неуплата штрафа за ДС (15 дней)")


# --- Генерация документов и E-mail (Заглушки) ---

def get_increase_penalty_doc(deal_id):
    """
    Генерирует уведомление о штрафе (10% + доплата).
    """
    deal_status = DealStatus.query.get(deal_id)
    if not deal_status or not deal_status.area_increase_penalty_amount:
        return None

    financials = get_single_deal_details_for_workflow(deal_id)

    doc = docx.Document()
    doc.add_heading('УВЕДОМЛЕНИЕ О ШТРАФЕ', 0)
    p = doc.add_paragraph(f"Уважаемый(ая) {financials.get('client_name', 'Клиент')},")
    p.add_run(
        f"\n\nПо Вашей квартире №{financials.get('property_id', 'N/A')} зафиксировано увеличение площади, требующее подписания Дополнительного Соглашения (ДС).")
    p.add_run(f"\n\nВы не явились в офис в установленный 10-дневный срок для подписания ДС.")
    p.add_run(f"\nСумма доплаты за метры: {deal_status.area_increase_payment_amount or 0:.2f} у.е.")
    p.add_run(f"\nСумма штрафа (10% от договора): {(financials.get('deal_sum', 0) * 0.10):.2f} у.е.")
    p.add_run(f"\n\nИТОГО К ОПЛАТЕ (в течение 15 дней): ")
    p.add_run(f"{deal_status.area_increase_penalty_amount:.2f} у.е.").bold = True

    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer, financials.get('property_id', 'doc')


def _send_termination_email(deal_status, reason):
    """(Заглушка) Отправляет email о расторжении."""
    print("--- ОТПРАВКА EMAIL О РАСТОРЖЕНИИ ---")
    print(f"ID Сделки: {deal_status.deal_id}")
    print(f"Причина: {reason}")
    # TODO: Добавить реальную логику отправки email
    # email_service.send_mail(...)
    print("---------------------------------------")